"""
Literacy code annotation engine.

Core approach (adapted from code_essay.py):
  - Model returns {"sentence": <verbatim>, "quote": <verbatim>, "code": <code>}
  - sentence anchors the quote so repeated words land in the right place
  - quote is verified verbatim against the paragraph before being tagged
  - anything that can't be located verbatim is dropped rather than guessed
  - tags are INSERTED into existing runs (non-destructive) so original
    formatting is preserved; coloured runs are then added for the quote

Color scheme (matches teacher examples):
  Orange  #E67E22  surface  : Sp, Caps, ^, WW, AA/Adj, Rep
  Red     #C0392B  grammar  : Gra, VT, V, WO, del, inc, RO, STR, Exp
  Blue    #2980B9  format   : P, FOR, //, Embed
  Green   #27AE60  positive : ✓  (positive Exp uses green)
"""

import json
import logging
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path.home() / ".writing_analyzer" / "coded"

# ── Literacy codes ─────────────────────────────────────────────────────────────
LITERACY_CODES: dict[str, tuple[str, str, bool]] = {
    "Sp":     ("Spelling error",             "E67E22", False),
    "Caps":   ("Missing capital letter",     "E67E22", False),
    "^":      ("Missing word",               "E67E22", False),
    "WW":     ("Wrong word",                 "E67E22", False),
    "AA/Adj": ("Wrong adjective form",       "E67E22", False),
    "Rep":    ("Repetition",                 "E67E22", False),
    "Gra":    ("Grammatical error",          "C0392B", False),
    "VT":     ("Verb tense error",           "C0392B", False),
    "V":      ("Verb formation error",       "C0392B", False),
    "WO":     ("Word order error",           "C0392B", False),
    "del":    ("Delete this word",           "C0392B", False),
    "inc":    ("Incomplete idea",            "C0392B", False),
    "RO":     ("Run-on sentence",            "C0392B", False),
    "STR":    ("Structure issue",            "C0392B", False),
    "Exp":    ("Expression issue",           "C0392B", False),   # error Exp = red
    "P":      ("Punctuation error",          "2980B9", False),
    "FOR":    ("Format problem",             "2980B9", False),
    "//":     ("New paragraph needed",       "2980B9", False),
    "Embed":  ("Quotation embedding error",  "2980B9", False),
    "✓":      ("Good point",                 "27AE60", True),    # positive only
}

VALID_CODES = set(LITERACY_CODES.keys())

# ── System prompt (sentence + quote format) ───────────────────────────────────
_CODEBOOK = """\
Sp     = spelling error ("recieved" → "received")
Caps   = missing capital letter ("i went" → "I went")
P      = punctuation missing/wrong (missing comma, apostrophe, hyphen in "so-called")
^      = a needed word is missing ("She ready" → "She is ready")
Exp    = awkward/unidiomatic expression that is not a clean grammar error
Gra    = grammatical error: subject-verb agreement, wrong/missing article, wrong preposition
Embed  = quotation embedded incorrectly (missing comma or quote marks)
AA/Adj = wrong adjective form ("most easiest" → "easiest")
STR    = sentence structure problem (clunky or ill-formed clause)
FOR    = formatting problem (register, heading, layout)
WO     = word order error ("To the store quickly went she")
WW     = wrong word, real word used incorrectly ("borrow" for "borrowed")
V      = missing or wrong verb formation ("They playing" → "They are playing")
VT     = verb tense error ("Yesterday he runs" → "ran")
del    = word should be deleted (redundant: "The dog, it barked")
inc    = incomplete sentence / fragment ("Because she was tired.")
RO     = run-on sentence (two clauses fused without punctuation)
Rep    = redundant repetition of a word or idea just used"""

_SYSTEM = f"""You are a precise English literacy marker. You find ERRORS in a student paragraph and label each with ONE code. You are conservative: only flag things that are clearly wrong.

CODES:
{_CODEBOOK}

RULES (follow exactly):
1. Only flag genuine errors. If a phrase is correct standard English, DO NOT flag it. When unsure, leave it alone.
2. "quote" must be copied VERBATIM from the paragraph, character for character. Never invent or paraphrase it.
3. "quote" must be whole words only. Never select part of a word (e.g. never "perfect" from "perfection").
4. "quote" is the SHORTEST span containing the error — usually one word; a short phrase only if the error spans multiple words.
5. Pick exactly ONE code per finding.
6. "sentence" is the FULL verbatim sentence the quote sits in. Copy it exactly.
7. Do not flag style, tone, or things you would merely prefer differently. Errors only.
8. A paragraph with no errors returns []. That is a valid correct answer.
9. Do NOT flag: "and" as RO, gerunds (-ing words) used correctly, correctly used adjectives or nouns.

OUTPUT FORMAT — return ONLY a JSON array, nothing else:
{{"sentence": "<verbatim sentence>", "quote": "<verbatim error span>", "code": "<one code>"}}

EXAMPLE:
Paragraph: They is playing outside and she recieved the ball, the game was fun.
Answer:
[{{"sentence":"They is playing outside and she recieved the ball, the game was fun.","quote":"is","code":"Gra"}},{{"sentence":"They is playing outside and she recieved the ball, the game was fun.","quote":"recieved","code":"Sp"}}]"""

# /no_think is in the user message — required for Qwen3
_USER_TEMPLATE = "/no_think\nParagraph:\n{paragraph_text}\n\nAnswer:"


def build_prompt(paragraph_text: str, tokenizer) -> str:
    user = _USER_TEMPLATE.format(paragraph_text=paragraph_text[:3000])
    if tokenizer and hasattr(tokenizer, "apply_chat_template"):
        msgs = [
            {"role": "system", "content": _SYSTEM},
            {"role": "user",   "content": user},
        ]
        try:
            return tokenizer.apply_chat_template(msgs, tokenize=False, add_generation_prompt=True)
        except Exception:
            pass
    return f"{_SYSTEM}\n\n{user}"


# ── Response parsing ──────────────────────────────────────────────────────────

def parse_response(raw: str) -> list[tuple[str, str, str]]:
    """
    Returns list of (sentence, quote, code) tuples.
    Drops any item where sentence or quote is blank or code is unknown.
    """
    raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL)
    raw = raw.replace("```json", "").replace("```", "")
    start, end = raw.find("["), raw.rfind("]")
    if start < 0 or end < 0:
        logger.warning("No JSON array in coder response: %.300s", raw)
        return []
    try:
        items = json.loads(raw[start: end + 1])
    except json.JSONDecodeError as exc:
        logger.warning("JSON parse error in coder: %s", exc)
        return []
    out = []
    for it in items:
        if not isinstance(it, dict):
            continue
        s = it.get("sentence", "")
        q = it.get("quote", "")
        c = it.get("code", "")
        if isinstance(s, str) and isinstance(q, str) and c in VALID_CODES and q.strip():
            out.append((s.strip(), q.strip(), c))
    return out


# ── Locating ──────────────────────────────────────────────────────────────────
# Adapted from code_essay.py: whitespace-tolerant, word-boundary safe.

def _locate(haystack: str, needle: str) -> tuple[int, int] | None:
    """Return (start, end) of needle in haystack, tolerant of whitespace runs.
    Whole-word safe on both ends. Returns None if not found."""
    needle = needle.strip()
    if not needle:
        return None
    tokens = needle.split()
    pat = r"\s+".join(re.escape(t) for t in tokens)
    if needle[0].isalnum() or needle[0] == "'":
        pat = r"(?<!\w)" + pat
    if needle[-1].isalnum() or needle[-1] in ("'", '"'):
        pat = pat + r"(?!\w)"
    m = re.search(pat, haystack)
    return (m.start(), m.end()) if m else None


def _find_quote_span(para_text: str, sentence: str, quote: str) -> tuple[int, int] | None:
    """
    Find (start, end) of quote in para_text, scoped to sentence when possible.
    Adapted from code_essay.py find_quote_offset — returns span, not just end offset.
    """
    sent_span = _locate(para_text, sentence)
    if sent_span:
        ss, se = sent_span
        seg = para_text[ss:se]
        q = _locate(seg, quote)
        if q:
            return (ss + q[0], ss + q[1])
    # Fallback: whole paragraph
    return _locate(para_text, quote)


# ── DOCX annotation ───────────────────────────────────────────────────────────

def annotate_docx(
    source_path: Path,
    model,
    tokenizer,
    on_progress=None,
) -> tuple[Path, int]:
    """
    Annotate a DOCX paragraph by paragraph. Returns (output_path, annotation_count).
    """
    from mlx_lm import generate as mlx_generate

    source_path = Path(source_path)
    tmp_path = None

    if source_path.suffix.lower() == ".pdf":
        from text_extraction import extract_text
        text = extract_text(source_path)
        tmp_path = Path(tempfile.mktemp(suffix=".docx"))
        from docx import Document as _Doc
        d = _Doc()
        for line in text.split("\n"):
            d.add_paragraph(line)
        d.save(str(tmp_path))
        work_path = tmp_path
    else:
        work_path = source_path

    doc = Document(str(work_path))

    content_indices = [
        i for i, p in enumerate(doc.paragraphs)
        if p.text.strip() and len(p.text.split()) >= 2
    ]
    total = len(content_indices)
    total_annotations = 0

    for done, pi in enumerate(content_indices):
        if on_progress:
            on_progress(done, total)

        para = doc.paragraphs[pi]
        prompt = build_prompt(para.text.strip(), tokenizer)
        response = mlx_generate(
            model, tokenizer,
            prompt=prompt,
            max_tokens=1024,
            verbose=False,
            temp=0.0,          # greedy / deterministic — fewer hallucinations
        )
        findings = parse_response(response)
        if not findings:
            continue

        tagged = _tag_paragraph(para, findings)
        total_annotations += tagged

    if on_progress:
        on_progress(total, total)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    stem = source_path.stem
    if stem.endswith(" coded"):
        stem = stem[:-6]
    output_path = OUTPUT_DIR / f"{stem} coded.docx"
    doc.save(str(output_path))

    if tmp_path:
        tmp_path.unlink(missing_ok=True)

    return output_path, total_annotations


# ── Run-level tagging ─────────────────────────────────────────────────────────

def _tag_paragraph(para, findings: list[tuple[str, str, str]]) -> int:
    """
    Apply coloured [CODE] annotations to a paragraph.

    Strategy (from code_essay.py, extended with colour):
    1. Locate each (sentence, quote) verbatim — drop anything not found.
    2. Sort right-to-left so earlier insertions don't shift later offsets.
    3. For each span: split the containing run(s) to colour the quote and
       insert a bold coloured [CODE] tag immediately after it.

    Returns the number of annotations placed.
    """
    para_text = para.text
    placed: list[tuple[int, int, str]] = []   # (start, end, code)
    seen: set[tuple[int, int]] = set()

    for sentence, quote, code in findings:
        span = _find_quote_span(para_text, sentence, quote)
        if span is None:
            logger.debug("Dropped (not found): %r code=%s", quote, code)
            continue
        start, end = span
        if (start, end) in seen:
            continue
        seen.add((start, end))
        placed.append((start, end, code))

    # Right-to-left so character offsets stay valid as we mutate runs
    placed.sort(key=lambda x: x[0], reverse=True)

    for start, end, code in placed:
        _, hex_color, _ = LITERACY_CODES.get(code, ("?", "888888", False))
        _color_span_and_insert_tag(para, start, end, hex_color, f"[{code}]")

    return len(placed)


def _color_span_and_insert_tag(para, start: int, end: int, hex_color: str, tag: str):
    """
    Color para.text[start:end] and insert a bold colored `tag` immediately after.
    Works by splitting runs at `start` and `end`, then coloring the middle run(s).
    """
    # Split at `end` first (right-to-left), then at `start`
    _split_run_at(para, end)
    _split_run_at(para, start)

    # Now rebuild the run map and color the middle portion
    rgb = _rgb(hex_color)
    cur = 0
    in_span = False
    last_span_run = None

    for run in para.runs:
        rlen = len(run.text)
        run_start = cur
        run_end = cur + rlen

        if run_start >= start and run_end <= end:
            run.font.color.rgb = rgb
            in_span = True
            last_span_run = run
        elif in_span:
            # We've passed the span — insert the tag run before this run
            break

        cur += rlen

    # Insert the [CODE] run after the last span run
    if last_span_run is not None:
        tag_run = _insert_run_after(para, last_span_run, tag)
        tag_run.font.color.rgb = rgb
        tag_run.bold = True
        # Preserve base font
        if last_span_run.font.name:
            tag_run.font.name = last_span_run.font.name
        if last_span_run.font.size:
            tag_run.font.size = last_span_run.font.size


def _split_run_at(para, offset: int):
    """Split the run containing `offset` into two runs at that character position."""
    cur = 0
    for run in para.runs:
        rlen = len(run.text)
        if cur < offset < cur + rlen:
            local = offset - cur
            before = run.text[:local]
            after  = run.text[local:]
            run.text = before
            # Create a new run after this one with identical formatting
            new_run = _insert_run_after(para, run, after)
            _copy_run_format(run, new_run)
            return
        cur += rlen


def _insert_run_after(para, ref_run, text: str):
    """Insert a new run with `text` immediately after `ref_run` in the paragraph XML."""
    from docx.oxml import OxmlElement
    new_r = OxmlElement("w:r")
    new_rpr = OxmlElement("w:rPr")
    new_r.append(new_rpr)
    new_t = OxmlElement("w:t")
    new_t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    ref_run._r.addnext(new_r)
    # Wrap in a Run object so callers can set font properties
    from docx.text.run import Run
    return Run(new_r, para)


def _copy_run_format(src, dst):
    """Copy bold/italic/font name+size from src run to dst run."""
    if src.bold is not None:
        dst.bold = src.bold
    if src.italic is not None:
        dst.italic = src.italic
    try:
        if src.font.name:
            dst.font.name = src.font.name
        if src.font.size:
            dst.font.size = src.font.size
    except Exception:
        pass


def _rgb(hex_str: str) -> RGBColor:
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))
